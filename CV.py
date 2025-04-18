from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH 
import pyttsx3

def speak(text):
    pyttsx3.speak(text)
    

document = Document()


#profile picture
document.add_picture(
    'kbob.png', 
    width = Inches(1.5)
)

# name phone number and email details
name = input('What is your name? ')
speak('Hello' + name + 'How are you do today?')

speak('What is your phone number? ')
phone_number = input('What is your phone number? ')

speak('What is your email? ')
email = input('What is your email? ')

document.add_paragraph(
    name.capitalize() + ' | ' + phone_number + ' | ' + email
)

# about me
document.add_heading('About me')
document.add_paragraph(
    input('Tell about yourself ')
)

# work experience
document.add_heading('Work Experience')
p = document.add_paragraph()

company = input('Enter company ')
from_date = input('From date ')
to_date = input('To Date ')

p.add_run(company + ' ').bold = True

p.add_run(from_date + '-' + to_date + '\n').italic = True

experience_details = input(
    'Describe your experirnce at ' + company + ' '
)
p.add_run(experience_details)

# more experiences
while True:
    has_more_experiences = input(
        'Do you have more experiences? Yes or No ')
    
    if has_more_experiences.lower() == 'yes':
        company = input('Enter company ')
        from_date = input('From date ')
        to_date = input('To Date ')

        p.add_run('\n' + company + ' ').bold = True

        p.add_run(from_date + '-' + to_date + '\n').italsic = True

        experience_details = input(
            'Describe your experirnce at ' + company + ' '
        )
        p.add_run(experience_details)
    else:
        break

document.add_heading('Skills')
skill = input('Enter Skill ')
p = document.add_paragraph(skill)
p.style = 'List Bullet'

while True:
    has_more_skills = input(
        'Do you have more skills? Yes or No '
    )

    if has_more_skills.lower() == 'yes':
        skill = input('Enter Skill ')
        p = document.add_paragraph(skill)
        p.style = 'List Bullet'
    else:
        break

# footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = "CV generated using Amigos and Input QuickBooks course project"

# Set to center
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

document.save('cv.docx')